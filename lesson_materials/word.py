from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# === Часть 1: Создание нового Word-документа ===
def create_word_file(filename):
    # Создаем новый документ
    doc = Document()

    # Добавляем заголовок
    doc.add_heading('Отчет о продажах', level=1)

    # Добавляем абзац с текстом
    paragraph = doc.add_paragraph('Этот отчет содержит информацию о продажах за прошлый месяц. ')
    paragraph.add_run('Обратите внимание на ключевые показатели.').bold = True

    # Добавляем таблицу
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    # Заполняем заголовки таблицы
    headers = ['Продукт', 'Количество', 'Выручка']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # Заполняем данные таблицы
    data = [
        ['Товар А', '10', '1000 руб.'],
        ['Товар Б', '5', '500 руб.'],
        ['Товар В', '8', '800 руб.']
    ]

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, cell_data in enumerate(row_data):
            table.cell(row_idx, col_idx).text = cell_data

    # Добавляем текст с форматированием
    final_paragraph = doc.add_paragraph('Итоговая выручка: ')
    run = final_paragraph.add_run('2300 руб.')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 128, 0)  # Зеленый цвет
    final_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Сохраняем документ
    doc.save(filename)
    print(f"Файл {filename} успешно создан.")


# === Часть 2: Чтение содержимого Word-документа ===
def read_word_file(filename):
    doc = Document(filename)
    print(f"Содержимое файла {filename}:\n")

    # Читаем текст из всех абзацев
    for paragraph in doc.paragraphs:
        print(paragraph.text)

    # Читаем данные из таблиц
    for table in doc.tables:
        print("\nДанные из таблицы:")
        for row in table.rows:
            print([cell.text for cell in row.cells])


# === Основной код ===
if __name__ == "__main__":
    file_name = "example.docx"

    # Создаем новый Word-файл
    create_word_file(file_name)

    # Читаем содержимое созданного файла
    read_word_file(file_name)
