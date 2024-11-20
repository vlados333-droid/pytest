from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_word_file(filename):
    doc = Document()

    doc.add_heading('Отчёт о продажах', level=1)

    paragraph = doc.add_paragraph('Этот отчёт содержит информацию о продажах за прошлый месяц. ')
    paragraph.add_run('Обратите внимание на ключевые показатели. ').bold = True

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    headers = ['Продукт', 'Количество', 'Выручка']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    data = [
        ['Машинки', '10', '1000 тнг'],
        ['Яблоки', '5 кг', '500 тнг'],
        ['Тетради', '100', '50 тнг'],
    ]

    for row_id, row_data in enumerate(data, start=1):
        for col_id, cell_data in enumerate(row_data):
            table.cell(row_id, col_id).text = cell_data

    final_paragraph = doc.add_paragraph('Итоговая выручка: ')
    run = final_paragraph.add_run('17 500 тнг')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 128, 0)
    final_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.save(filename)
    print(f"Файл {filename} успешно создан")


def read_word_file(filename):
    doc = Document(filename)
    print(f"Содержимое файла {filename}: \n")

    for paragraph in doc.paragraphs:
        print(paragraph.text)

    for table in doc.tables:
        print("\nДанные из таблицы: ")
        for row in table.rows:
            print([cell.text for cell in row.cells])


if __name__ == "__main__":
    file_name = 'example.docx'

    create_word_file(file_name)

    read_word_file(file_name)
